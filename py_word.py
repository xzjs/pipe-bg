from docx import Document
from docx.shared import Inches
from flask import Flask, request

app = Flask(__name__)

def replace_text_img(doc, old_word, new_word):
    for paragraph in doc.paragraphs:
        # 替换段落中的参数
        if old_word in paragraph.text:
            paragraph.text = paragraph.text.replace(old_word, new_word)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                text = cell.text.strip()  # 获取单元格的文本内容

                # 判断文本内容是否为参数名，并替换为相应的参数值
                if text == old_word:
                    value = new_word

                    # 判断是否是图像参数
                    if value.endswith(('.png', '.jpg')):
                        # 替换为图片
                        cell.paragraphs[0].clear()
                        cell.paragraphs[0].add_run().add_picture(value, width=Inches(2))
                    else:
                        # 替换为文本
                        cell.text = str(value)


@app.route('/api/report', methods=['POST'])
def replace_context():
    doc = Document('doc.docx')

    data = request.get_json()
    no = data['no']
    method = data['method']
    file = data['file']
    start_no = data['start_no']
    end_no = data['end_no']
    time = data['time']
    start_depth = data['start_depth']
    end_depth = data['end_depth']
    type = data['type']
    material = data['material']
    diameter = data['diameter']
    direction = data['direction']

    replace_text_img(doc, 'no', no)
    replace_text_img(doc, 'method', method)
    replace_text_img(doc, 'file', file)
    replace_text_img(doc, 'start_no', start_no)
    replace_text_img(doc, 'end_no', end_no)
    replace_text_img(doc, 'time', time)
    replace_text_img(doc, 'start_depth', start_depth)
    replace_text_img(doc, 'end_depth', end_depth)
    replace_text_img(doc, 'type', type)
    replace_text_img(doc, 'material', material)
    replace_text_img(doc, 'diameter', diameter)
    replace_text_img(doc, 'direction', direction)

    doc.save('new_Doc.docx')

    return {"msg": 'ok'}


if __name__ == '__main__':
    app.run()


# def replace_context(docx):
#     # post
#     # filename=post.request.
#     id_no = '0'
#     met = '0'
#     name = '0.mp4'
#     start_0 = '2'
#     end_0 = '4.5'
#     date = '2020年1月2日'
#     depth_0 = '1'
#     depth_1 = '10'
#     type = '通风管道'
#     mater = '塑料'
#     diame = '10'
#     direc = '水平'
#     time = '6:00'
#     img_path1 = './img/0.jpg'
#     img_path2 = './img/0.jpg'
#     img_path3 = './img/0.jpg'
#     img_path4 = './img/0.jpg'
#
#     replace_text_img(doc, 'no', id_no)
#     replace_text_img(doc, 'method', met)
#     replace_text_img(docx, 'file', name)
#     replace_text_img(docx, 'start_no', start_0)
#     replace_text_img(docx, 'end_no', end_0)
#     replace_text_img(docx, 'time', date)
#     replace_text_img(docx, 'start_depth', depth_0)
#     replace_text_img(docx, 'end_depth', depth_1)
#     replace_text_img(docx, 'type', type)
#     replace_text_img(docx, 'material', mater)
#     replace_text_img(docx, 'diameter', diame)
#     replace_text_img(docx, 'direction', direc)
#     replace_text_img(docx, 'time1', time)
#     replace_text_img(docx, 'img1', img_path1)
#     replace_text_img(docx, 'img2', img_path2)
#     replace_text_img(docx, 'img3', img_path3)
#     replace_text_img(docx, 'img4', img_path4)
#
#     doc.save('new_Doc.docx')