import json
import os
import time
from flask import Flask, abort, jsonify, redirect, request, url_for
import ffmpeg
from sqlalchemy import create_engine, select
from models import Mark, Video, Img, Base
import cv2
from sqlalchemy.orm import Session
from docx import Document
from docx.shared import Inches

ALLOWED_EXTENSIONS = {'mp4'}
app = Flask(__name__)
app.config.from_file("config.json", load=json.load)
SQLALCHEMY_DATABASE_URI = f"mysql+pymysql://{app.config['USERNAME']}:{app.config['PWD']}@{app.config['HOST']}:{app.config['PORT']}/{app.config['DATABASE']}?charset=utf8mb4"
engine = create_engine(SQLALCHEMY_DATABASE_URI)
Base.metadata.create_all(bind=engine)


@app.route("/api/extract", methods=['POST'])
def extract():
    data = request.get_json()
    video_id = data['video_id']
    with Session(engine) as session:
        video = session.scalars(
            select(Video).where(Video.id == video_id)).first()
        path = app.config['SAVE_FOLDER'] + video.src
        current_time = data['current_time']
        probe = ffmpeg.probe(path)
        video_info = next(
            s for s in probe['streams'] if s['codec_type'] == 'video')
        width = int(video_info['width'])
        height = int(video_info['height'])
        r_frame_rate = probe['streams'][0]['r_frame_rate'].split('/')
        frame_rate = int(r_frame_rate[0])/int(r_frame_rate[1])
        frame_num = int(current_time*frame_rate)
        output = "/%d.jpg" % time.time()
        (
            ffmpeg
            .input(path)
            .filter('select', 'gte(n,{})'.format(frame_num))
            .output(app.config['SAVE_FOLDER']+output, vframes=1, format='image2', vcodec='mjpeg')
            .run()
        )
        img = Img(output, video_id, width=width, height=height)
        session.add(img)
        session.commit()
        return jsonify(img.to_dict())


def allowed_file(filename):
    return '.' in filename and \
           filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS


@app.route("/api/upload", methods=['POST'])
def upload_file():
    if 'file' in request.files:
        file = request.files['file']
        if file.filename and allowed_file(file.filename):
            filename = '/%d.mp4' % time.time()
            save_path = app.config['SAVE_FOLDER'] + filename
            file.save(save_path)
            video = Video(src=filename)
            with Session(engine) as session:
                session.add(video)
                session.commit()
                return jsonify(video.to_dict())
    abort(400)


@app.route("/api/mark", methods=['POST'])
def mark_post():
    data = request.get_json()
    src = data['src']
    flaw = data['flaw']
    level = data['level']
    x = data['x']
    y = data['y']
    width = data['width']
    height = data['height']
    img_id = data['img_id']

    with Session(engine) as session:
        imgObj = session.scalars(select(Img).where(Img.id == img_id)).first()
        img = cv2.imread(app.config['SAVE_FOLDER'] + imgObj.src)
        pt1 = (x, y)
        pt2 = (x+width, y+height)
        cv2.rectangle(img, pt1, pt2, (0, 0, 255), 1)
        name = "/%d.jpg" % time.time()
        cv2.imwrite(app.config['SAVE_FOLDER'] + name, img)
        mark = Mark(name, flaw, level, x, y, width, height, img_id)
        session.add(mark)
        session.commit()
        return jsonify(mark.to_dict())


@app.route("/api/video/<id>", methods=['GET'])
def video_get(id):
    with Session(engine) as session:
        video = session.scalars(select(Video).where(Video.id == id)).first()
        return jsonify(video.to_dict())


@app.route("/api/report", methods=['POST'])
def report_post():
    data = request.get_json()
    report_keys = ['no', 'method', 'file', 'start_no', 'end_no', 'time', 'start_depth', 'end_depth', 'type', 'material', 'diameter', 'direction',
                   'pipe_length', 'detection_length', 'repair_index', 'maintenance_index', 'inspection_personnel', 'detection_site', 'detection_date', 'remark']
    detail_keys = ['distance', 'flaw', 'score', 'level',
                   'description', 'description', 'position', 'img_id']  # description占了两列
    img_position = [(12, 0), (12, 5), (14, 0), (14, 5)]
    name_position = [(13, 0), (13, 5), (15, 0), (15, 5)]
    details = data['detail']
    doc = Document(app.config['SAVE_FOLDER'] + '/doc.docx')
    for key in report_keys:
        for paragraph in doc.paragraphs:
            paragraph.text = paragraph.text.replace(key, str(data[key]))
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text = cell.text.strip()
                    if text == key:
                        cell.text = data[key]

    for table in doc.tables:
        base_index = 7
        for i in range(len(details)):
            detail = details[i]
            for j in range(len(detail_keys)):
                cell = table.cell(base_index+i, j)
                key = detail_keys[j]
                cell.text = str(detail[key])
            position = img_position[i]
            cell = table.cell(position[0], position[1])
            cell.paragraphs[0].clear()
            cell.paragraphs[0].add_run().add_picture(
                app.config['SAVE_FOLDER']+detail['src'], width=Inches(2))
            position = name_position[i]
            cell = table.cell(position[0], position[1])
            cell.paragraphs[0].text ='图片'+str(detail['img_id'])

    name = '/%d.docx' % time.time()
    doc.save(app.config['SAVE_FOLDER']+name)
    return jsonify({'src': name})


def replace_text_img(doc, old_word, new_word):
    for paragraph in doc.paragraphs:
        # 替换段落中的参数
        if old_word in paragraph.text:
            paragraph.text = paragraph.text.replace(old_word, new_word)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                text = cell.text.strip()  # 获取单元格的文本内容

                # 判断文本内容是否为参数名，并替换为相应的参数值
                if text == old_word:
                    value = new_word

                    # 判断是否是图像参数
                    if value.endswith(('.png', '.jpg')):
                        # 替换为图片
                        cell.paragraphs[0].clear()
                        cell.paragraphs[0].add_run().add_picture(
                            value, width=Inches(2))
                    else:
                        # 替换为文本
                        cell.text = str(value)
